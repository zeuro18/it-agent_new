"""
ingest.py
Document ingestion pipeline for the RAG system.

Reads PDF and DOCX policy documents from docs/, chunks them,
embeds with sentence-transformers, and stores in ChromaDB + BM25 index.
"""

import os
import re
import json
import pickle
from pathlib import Path

from sentence_transformers import SentenceTransformer
import chromadb
from rank_bm25 import BM25Okapi
from pypdf import PdfReader
import docx

DOCS_DIR = os.path.join(os.path.dirname(__file__), "docs")
CHROMA_DIR = os.path.join(os.path.dirname(__file__), "chroma_store")
BM25_PATH = os.path.join(os.path.dirname(__file__), "bm25_index.pkl")
CHUNKS_PATH = os.path.join(os.path.dirname(__file__), "chunks.json")

EMBED_MODEL = "all-MiniLM-L6-v2"
CHUNK_SIZE = 400       # characters per chunk
CHUNK_OVERLAP = 80     # character overlap between chunks


def chunk_text(text: str, source: str, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP) -> list[dict]:
    """
    Split a document into overlapping chunks.
    Each chunk carries metadata about its source file and position.
    """
    chunks = []
    # Split by section headers first for better semantic boundaries
    sections = re.split(r'\n(?=##?\s)', text)

    chunk_id = 0
    for section in sections:
        section = section.strip()
        if not section:
            continue

        # Extract section title if present
        title_match = re.match(r'^(##?\s+.+)', section)
        section_title = title_match.group(1).strip('# ').strip() if title_match else ""

        # If section is small enough, keep it whole
        if len(section) <= chunk_size:
            chunks.append({
                "id": f"{source}__chunk_{chunk_id}",
                "text": section,
                "source": source,
                "section": section_title,
            })
            chunk_id += 1
        else:
            # Slide a window over the section
            start = 0
            while start < len(section):
                end = start + chunk_size
                chunk_text_slice = section[start:end]
                chunks.append({
                    "id": f"{source}__chunk_{chunk_id}",
                    "text": chunk_text_slice,
                    "source": source,
                    "section": section_title,
                })
                chunk_id += 1
                start += chunk_size - overlap

    return chunks


def ingest():
    """
    Full ingestion pipeline:
    1. Read all .pdf and .docx files from docs/
    2. Chunk them
    3. Embed with sentence-transformers → store in ChromaDB
    4. Build BM25 index over the same chunks
    """
    print("[ingest] Loading documents from", DOCS_DIR)

    all_chunks = []
    for fname in os.listdir(DOCS_DIR):
        fpath = os.path.join(DOCS_DIR, fname)
        text = ""
        
        if fname.endswith(".pdf"):
            print(f"[ingest] Reading PDF: {fname}")
            try:
                reader = PdfReader(fpath)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            except Exception as e:
                print(f"Error reading PDF {fname}: {e}")
                
        elif fname.endswith(".docx"):
            print(f"[ingest] Reading DOCX: {fname}")
            try:
                doc = docx.Document(fpath)
                # Extract paragraph text
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
                # Also extract table cell text (CIS templates store requirements in tables)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = "  |  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            text += row_text + "\n"
            except Exception as e:
                print(f"Error reading DOCX {fname}: {e}")
        else:
            continue
            
        if text.strip():
            chunks = chunk_text(text, source=fname)
            all_chunks.extend(chunks)
            print(f"  [{fname}] -> {len(chunks)} chunks")

    if not all_chunks:
        print("[ingest] No documents found!")
        return

    print(f"[ingest] Total chunks: {len(all_chunks)}")

    # save chunk metadata for later retrieval
    with open(CHUNKS_PATH, "w", encoding="utf-8") as f:
        json.dump(all_chunks, f, indent=2)

    #Dense embeddings via sentence-transformers → ChromaDB
    print(f"[ingest] Loading embedding model: {EMBED_MODEL}")
    model = SentenceTransformer(EMBED_MODEL)

    texts = [c["text"] for c in all_chunks]
    ids = [c["id"] for c in all_chunks]
    metadatas = [{"source": c["source"], "section": c["section"]} for c in all_chunks]

    print("[ingest] Generating embeddings...")
    embeddings = model.encode(texts, show_progress_bar=True).tolist()

    print("[ingest] Storing in ChromaDB...")
    client = chromadb.PersistentClient(path=CHROMA_DIR)
    # Delete existing collection if it exists to avoid duplicates
    try:
        client.delete_collection("it_policies")
    except Exception:
        pass
    collection = client.create_collection("it_policies")
    collection.add(ids=ids, embeddings=embeddings, documents=texts, metadatas=metadatas)

    # BM25 index 
    print("[ingest] Building BM25 index...")
    tokenized = [t.lower().split() for t in texts]
    bm25 = BM25Okapi(tokenized)
    with open(BM25_PATH, "wb") as f:
        pickle.dump({"bm25": bm25, "ids": ids, "texts": texts, "metadatas": metadatas}, f)

    print(f"[ingest] Done! {len(all_chunks)} chunks indexed.")
    print(f"  ChromaDB: {CHROMA_DIR}")
    print(f"  BM25:     {BM25_PATH}")


if __name__ == "__main__":
    ingest()
